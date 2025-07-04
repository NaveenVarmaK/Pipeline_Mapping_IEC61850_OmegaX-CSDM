#
# Final Enhanced Script to Extract Data from Word Document Tables
#
# REVISION 5:
# - Added a de-duplication step to prevent processing the same Data Object multiple times.
# - The script now tracks processed 'Data Object (DO)' names and skips any duplicates.
#
# FEATURES:
# - Automatic "Data Object" table recognition.
# - De-duplication of data objects before LLM processing.
# - CPU and memory usage monitoring.
# - Detailed timing and performance reporting.
#
# Pre-requisites:
# 1. Run `pip install python-docx openai psutil lxml`
# 2. Have LM Studio running with a model loaded, serving on http://localhost:1234
# 3. An editable CSV_Header_Dictionary.py file.
#

import docx
import json
import openai
import argparse
import os
import sys
import time
import psutil
import threading
from typing import List, Dict, Optional, Generator
from docx.table import Table

# --- Configuration (remains the same) ---
DEFAULT_WORD_FILE = "C:/Users/Naveen/Documents/Internship_EMSE2/OmegaX-Pipeline/Resources/Library_EDF_OmegaX_rev02_29june2023.docx"
DEFAULT_LLM_URL = "http://localhost:1234/v1"
DEFAULT_CSV_DICT_FILE = "Resources/CSV_Header_Dictionary.py"
DATA_OBJECT_KEYWORDS = ["data object", "dataobject", "data objects description"]
DATA_OBJECT_HEADERS = ["data object (do)", "type", "description", "m/o in standard"]

# --- Prompt (remains the same) ---
PROMPT_TEMPLATE = """
You are an expert ontology engineer. Your task is to convert a description of a data object into a structured JSON format. The output must be a single, valid JSON object and nothing else.

QUDT_UNIT_EXAMPLES = 
- watts (W): W
- active power (W): W
- kilowatts (kW): KiloW
- megawatts (MW): MegaW
- volt-amperes (VA): V-A
- kilo volt-amperes (kVA): KiloV-A
- mega volt-amperes (MVA): MegaV-A
- volt-amperes reactive (VAr): V-A_Reactive
- kilo volt-amperes reactive (kVAr): KiloV-A_Reactive
- mega volt-amperes reactive (MVAr): MegaV-A_Reactive
- watt-hours (Wh): W-HR
- kilowatt-hours (kWh): KiloW-R
- megawatt-hours (MWh): MegaW-R

- volts (V): V
- kilovolts (kV): KiloV
- amperes (A): A
- kiloamperes (kA): KiloA
- ohms (ohm): OHM
- hertz (Hz): HZ
- power factor (no unit): UNITLESS 
- degrees (phase angle): DEG

- meters per second (m/s): M-Per-SEC
- watts per square meter (W/m²): W-M2
- degrees Celsius (degC): DEG_C
- kelvin (K): K
- pascals (Pa): PA
- hectopascals (hPa): HectoPA
- percentage (%): PERCENT

- seconds (s): SEC
- minutes (min): MIN
- hours (h): HR

- meters (m): M
- millimeters (mm): MilliM
- square meters (m²): M2
- cubic meters (m³): M3
- No specific unit: UNITLESS

you must use the QUDT unit examples to determine the correct unit for the data object.

Here are examples of the conversion:

## EXAMPLE 1 ##
Input:
{{
  "Data Object (DO)": "SnwFll",
  "Type": "MV-STD_MV",
  "Description": "Snowfall (typically in mm length SIUnit [m]).",
  "Description du standard": "Snowfall (typically in mm length SIUnit [m]).",
  "M/O in standard": "O"
}}
Output:
{{
    "snwfll": {{
        "description": "Snowfall (typically in mm - length SIUnit [m])",
        "unit": "MilliM",
        "enum_kind": "MV_STD",
        "property": "Snowfall"
    }}
}}

## EXAMPLE 2 ##
Input:
{{
  "Data Object (DO)": "RisNgGnd",
  "Type": "MV-STD_MV",
  "Description": "DC resistance between negative pole and earth",
  "Description du standard": "DC resistance between negative pole and earth",
  "M/O in standard": "O"
}}
Output:
{{
    "risnggnd": {{
        "description": "DC resistance between negative pole and earth",
        "unit": "OHM",
        "type": "MV-STD_MV",
        "property": "DCResistanceNegativePoleToEarth"
    }}
}}

---
Now, perform the same conversion for the following new data object. Provide ONLY the JSON object as your output.

## NEW DATA OBJECT ##
Input:
{input_data}
Output:
"""


# --- Monitoring Classes (remain the same) ---
class ResourceMonitor:
    """Monitor CPU and memory usage during processing."""

    def __init__(self):
        self.monitoring = False
        self.cpu_samples = []
        self.memory_samples = []
        self.monitor_thread = None
        self.process = psutil.Process()

    def start_monitoring(self):
        self.monitoring = True
        self.monitor_thread = threading.Thread(target=self._monitor_resources, daemon=True)
        self.monitor_thread.start()

    def stop_monitoring(self):
        self.monitoring = False
        if self.monitor_thread:
            self.monitor_thread.join(timeout=2)

    def _monitor_resources(self):
        while self.monitoring:
            try:
                cpu_percent = self.process.cpu_percent(interval=1.0)
                memory_mb = self.process.memory_info().rss / (1024 * 1024)
                self.cpu_samples.append(cpu_percent)
                self.memory_samples.append(memory_mb)
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                break

    def get_statistics(self):
        if not self.cpu_samples or not self.memory_samples:
            return {'cpu_avg': 0, 'cpu_max': 0, 'mem_avg_mb': 0, 'mem_peak_mb': 0}
        return {
            'cpu_avg': sum(self.cpu_samples) / len(self.cpu_samples),
            'cpu_max': max(self.cpu_samples),
            'mem_avg_mb': sum(self.memory_samples) / len(self.memory_samples),
            'mem_peak_mb': max(self.memory_samples)
        }


class PerformanceTracker:
    """Track timing and performance metrics."""

    def __init__(self):
        self.start_time = 0.0
        self.end_time = 0.0
        self.llm_times = []
        self.extraction_time = 0.0
        self.update_time = 0.0

    def start_total_timing(self):
        self.start_time = time.monotonic()

    def end_total_timing(self):
        self.end_time = time.monotonic()

    def add_llm_time(self, duration):
        self.llm_times.append(duration)

    def set_extraction_time(self, duration):
        self.extraction_time = duration

    def set_update_time(self, duration):
        self.update_time = duration

    def get_statistics(self):
        total_time = (self.end_time - self.start_time) if self.start_time and self.end_time else 0
        llm_stats = {}
        if self.llm_times:
            llm_stats = {
                'total_llm_time': sum(self.llm_times),
                'avg_llm_time': sum(self.llm_times) / len(self.llm_times),
                'max_llm_time': max(self.llm_times),
                'min_llm_time': min(self.llm_times),
                'llm_calls': len(self.llm_times)
            }
        return {'total_time': total_time, 'extraction_time': self.extraction_time, 'update_time': self.update_time,
                **llm_stats}


class WordDocumentProcessor:
    """Processor combining your original logic with advanced monitoring and auto-detection."""

    def __init__(self, word_file_path, llm_base_url, csv_dict_file, dry_run=False):
        self.word_file_path = word_file_path
        self.csv_dict_file = csv_dict_file
        self.dry_run = dry_run
        self.client = openai.OpenAI(base_url=llm_base_url, api_key="not-needed")
        self.processed_results = []
        self.resource_monitor = ResourceMonitor()
        self.performance_tracker = PerformanceTracker()
        self._validate_or_prepare_dictionary_file()

    def _validate_or_prepare_dictionary_file(self):
        if self.dry_run:
            return
        if os.path.exists(self.csv_dict_file):
            print(f"--- Validating existing dictionary file: {self.csv_dict_file} ---")
            with open(self.csv_dict_file, 'r', encoding='utf-8') as f:
                content = f.read()
            if "MEASUREMENTS = {" not in content:
                print(f"\n[ERROR] FATAL: The file '{self.csv_dict_file}' is invalid!", file=sys.stderr)
                print("It exists, but does not contain the required 'MEASUREMENTS = {' dictionary definition.",
                      file=sys.stderr)
                print("Please correct the file or delete it to allow automatic creation.", file=sys.stderr)
                sys.exit(1)
            print("  [✓] File is valid.")
        else:
            print(f"--- [INFO] Dictionary file '{self.csv_dict_file}' not found. ---")
            print("--- It will be created automatically upon successful processing. ---")

    def _is_data_object_table(self, table: Table, preceding_text: str = "") -> bool:
        if preceding_text:
            if any(keyword in preceding_text for keyword in DATA_OBJECT_KEYWORDS):
                return True
        if len(table.rows) < 1: return False
        for row_idx in range(min(2, len(table.rows))):
            try:
                headers = [cell.text.strip().lower() for cell in table.rows[row_idx].cells]
                if sum(1 for expected in DATA_OBJECT_HEADERS if any(expected in h for h in headers)) >= 2:
                    return True
            except IndexError:
                continue
        return False

    def identify_data_object_tables(self) -> List[int]:
        try:
            doc = docx.Document(self.word_file_path)
            indices = []
            print("\n--- Scanning document for Data Object tables ---")
            table_num = 0
            for block in doc.element.body:
                if block.tag.endswith('tbl'):
                    preceding_text = ""
                    prev_sibling = block.getprevious()
                    if prev_sibling is not None and prev_sibling.tag.endswith('p'):
                        p_text = "".join(node.text for node in prev_sibling.xpath('.//w:t'))
                        preceding_text = p_text.strip().lower()
                    if self._is_data_object_table(doc.tables[table_num], preceding_text):
                        indices.append(table_num)
                        print(f"  [✓] Table {table_num}: Identified as a Data Object table.")
                    else:
                        print(f"  [✗] Table {table_num}: Skipping.")
                    table_num += 1
            if not indices:
                print("\nWarning: No Data Object tables were identified.")
            else:
                print(f"\nFound {len(indices)} table(s) to process: {indices}")
            return indices
        except Exception as e:
            print(f"Error during table identification: {e}", file=sys.stderr)
            if "lxml" in str(e): print("Hint: Run 'pip install lxml'.", file=sys.stderr)
            return []

    def extract_data_from_word(self, table_indices: List[int]):
        start_time = time.monotonic()
        try:
            doc = docx.Document(self.word_file_path)
            for table_idx in table_indices:
                print(f"\n--- Extracting data from Table {table_idx} ---")
                table = doc.tables[table_idx]
                if len(table.rows) < 2:
                    print(f"  Warning: Table {table_idx} has fewer than two rows. Skipping.")
                    continue
                headers = [cell.text.strip() for cell in table.rows[1].cells]
                print(f"  Headers found: {headers}")
                for row in table.rows[2:]:
                    row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells) if i < len(headers)}
                    if row_data.get("Data Object (DO)"):
                        yield row_data
        except Exception as e:
            print(f"Error extracting data: {e}", file=sys.stderr)
        finally:
            self.performance_tracker.set_extraction_time(time.monotonic() - start_time)

    def get_structured_data_from_llm(self, data_from_word):
        start_time = time.monotonic()
        input_str = json.dumps(data_from_word, indent=2)
        prompt = PROMPT_TEMPLATE.format(input_data=input_str)
        try:
            completion = self.client.chat.completions.create(
                model="local-model",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
            )
            response_content = completion.choices[0].message.content.strip()
            if response_content.startswith("```json"): response_content = response_content[7:]
            if response_content.endswith("```"): response_content = response_content[:-3]
            return json.loads(response_content.strip())
        except Exception as e:
            print(f"LLM Error: {e}", file=sys.stderr)
            return None
        finally:
            self.performance_tracker.add_llm_time(time.monotonic() - start_time)

    def update_csv_dictionary(self):
        if self.dry_run or not self.processed_results:
            if self.dry_run: print("\n--- [DRY RUN] Skipping dictionary update. ---")
            return
        start_time = time.monotonic()
        try:
            if not os.path.exists(self.csv_dict_file):
                with open(self.csv_dict_file, 'w', encoding='utf-8') as f:
                    f.write("# Auto-generated dictionary\n\nMEASUREMENTS = {\n}\n")
            with open(self.csv_dict_file, 'r', encoding='utf-8') as f:
                content = f.read()
            measurements_start = content.find("MEASUREMENTS = {")
            brace_count = 0
            measurements_end = measurements_start + len("MEASUREMENTS = {")
            for i, char in enumerate(content[measurements_end:], measurements_end):
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    if brace_count == 0:
                        measurements_end = i
                        break
                    brace_count -= 1
            measurements_content = content[measurements_start + len("MEASUREMENTS = {"):measurements_end]
            new_entries = [f'    "{key}": {json.dumps(value, indent=8)}' for d in self.processed_results for key, value
                           in d.items()]
            if measurements_content.strip() and not measurements_content.strip().endswith(','):
                measurements_content += ','
            new_content = measurements_content + '\n' + ',\n'.join(new_entries)
            updated_content = content[:measurements_start + len("MEASUREMENTS = {")] + new_content + content[
                                                                                                     measurements_end:]
            with open(self.csv_dict_file, 'w', encoding='utf-8') as f:
                f.write(updated_content)
            print(f"\nSuccessfully updated {self.csv_dict_file} with {len(new_entries)} new entries.")
        except Exception as e:
            print(f"Error updating dictionary file: {e}", file=sys.stderr)
        finally:
            self.performance_tracker.set_update_time(time.monotonic() - start_time)

    def print_performance_report(self):
        print("\n" + "=" * 80)
        print("PERFORMANCE AND RESOURCE REPORT".center(80))
        print("=" * 80)
        timing = self.performance_tracker.get_statistics()
        resource = self.resource_monitor.get_statistics()
        print(f"\n⏱️  TIMING STATISTICS:")
        print(f"   - Total Processing Time: {timing.get('total_time', 0):.2f}s")
        print(f"   - Data Extraction (Word): {timing.get('extraction_time', 0):.2f}s")
        if not self.dry_run: print(f"   - Dictionary File Update: {timing.get('update_time', 0):.2f}s")
        if 'llm_calls' in timing:
            print(f"\n🤖 LLM PERFORMANCE ({timing['llm_calls']} calls):")
            print(f"   - Avg Response Time: {timing['avg_llm_time']:.2f} s/prompt")
            print(f"   - Total LLM Time:    {timing['total_llm_time']:.2f}s")
        print(f"\n💻 RESOURCE USAGE:")
        print(f"   - Average CPU: {resource.get('cpu_avg', 0):.1f}% | Peak CPU: {resource.get('cpu_max', 0):.1f}%")
        print(
            f"   - Average Mem: {resource.get('mem_avg_mb', 0):.1f} MB | Peak Mem: {resource.get('mem_peak_mb', 0):.1f} MB")
        print("=" * 80)

    # UPDATED: This is the main change to handle de-duplication.
    def process_document(self):
        """Main workflow combining all features, now with de-duplication."""
        print("--- Starting Enhanced Data Processing ---")
        self.performance_tracker.start_total_timing()
        self.resource_monitor.start_monitoring()
        try:
            table_indices = self.identify_data_object_tables()
            if not table_indices:
                return

            # NEW: Keep track of processed items and counts for the summary report.
            processed_do_names = set()
            processed_count = 0
            skipped_duplicates_count = 0
            failed_count = 0

            data_generator = self.extract_data_from_word(table_indices)

            for item_from_word in data_generator:
                do_name = item_from_word.get("Data Object (DO)")

                # NEW: De-duplication check.
                if do_name in processed_do_names:
                    print(f"  - Skipping duplicate Data Object: '{do_name}'")
                    skipped_duplicates_count += 1
                    continue  # Move to the next item

                # If it's not a duplicate, process it.
                print("-" * 60)
                print(f"Processing item {processed_count + 1}: '{do_name}'")

                # Add the name to the set to mark it as processed.
                processed_do_names.add(do_name)

                structured_result = self.get_structured_data_from_llm(item_from_word)
                if structured_result:
                    print("  [✓] LLM structuring successful.")
                    self.processed_results.append(structured_result)
                    processed_count += 1
                else:
                    print("  [✗] LLM structuring failed.")
                    failed_count += 1

            self.update_csv_dictionary()

            # UPDATED: The summary report now includes skipped and failed counts.
            print("\n--- Processing Summary ---")
            print(f"  - Successfully processed: {processed_count} unique items")
            print(f"  - Skipped duplicates:     {skipped_duplicates_count} items")
            print(f"  - Failed to process:      {failed_count} items")

        finally:
            self.performance_tracker.end_total_timing()
            self.resource_monitor.stop_monitoring()
            self.print_performance_report()


def parse_arguments():
    """Parse command line arguments. Note: --table-index is no longer needed."""
    parser = argparse.ArgumentParser(description="Auto-detect and process data object tables from a Word doc.")
    parser.add_argument("--word-file", "-w", default=DEFAULT_WORD_FILE, help="Path to the Word document.")
    parser.add_argument("--llm-url", "-l", default=DEFAULT_LLM_URL, help="LLM server URL.")
    parser.add_argument("--csv-dict", "-c", default=DEFAULT_CSV_DICT_FILE,
                        help="Path to the Python dictionary file to update.")
    parser.add_argument("--dry-run", "-d", action="store_true", help="Run without modifying files.")
    return parser.parse_args()


def main():
    args = parse_arguments()
    processor = WordDocumentProcessor(
        word_file_path=args.word_file,
        llm_base_url=args.llm_url,
        csv_dict_file=args.csv_dict,
        dry_run=args.dry_run
    )
    processor.process_document()


if __name__ == "__main__":
    main()